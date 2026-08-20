import re
from dataclasses import dataclass
from pathlib import Path

import pymupdf
from docx import Document as WordDocument


class DocumentParseError(ValueError):
    """文档内容损坏、加密或无法提取文本。"""


@dataclass(frozen=True, slots=True)
class ParsedPage:
    """解析器输出的原始页文本。"""

    page_number: int
    content: str


def normalize_text(text: str) -> str:
    """统一换行、清理行尾空格和过多空行。"""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = "\n".join(line.rstrip() for line in normalized.splitlines())
    return re.sub(r"\n{3,}", "\n\n", normalized).strip()


def ensure_has_text(pages: list[ParsedPage]) -> list[ParsedPage]:
    if not any(page.content for page in pages):
        raise DocumentParseError("文档中没有提取到文本，扫描版 PDF 暂不支持")
    return pages


def parse_pdf(path: Path) -> list[ParsedPage]:
    try:
        with pymupdf.open(path) as document:
            if document.needs_pass:
                raise DocumentParseError("加密 PDF 暂不支持")
            pages = [
                ParsedPage(page_number=index + 1, content=normalize_text(page.get_text("text")))
                for index, page in enumerate(document)
            ]
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError("PDF 文件损坏或无法解析") from exc
    return ensure_has_text(pages)


def parse_docx(path: Path) -> list[ParsedPage]:
    try:
        document = WordDocument(path)
        blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                blocks.append("\t".join(cell.text.strip() for cell in row.cells))
    except Exception as exc:
        raise DocumentParseError("DOCX 文件损坏或无法解析") from exc

    return ensure_has_text([ParsedPage(page_number=1, content=normalize_text("\n".join(blocks)))])


def read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("文本文件编码不受支持，请使用 UTF-8 或 GB18030")


def parse_text(path: Path) -> list[ParsedPage]:
    return ensure_has_text(
        [ParsedPage(page_number=1, content=normalize_text(read_text_file(path)))]
    )


def parse_document(path: str | Path) -> list[ParsedPage]:
    """根据扩展名选择解析器。"""
    document_path = Path(path)
    extension = document_path.suffix.lower()
    if extension == ".pdf":
        return parse_pdf(document_path)
    if extension == ".docx":
        return parse_docx(document_path)
    if extension in {".txt", ".md"}:
        return parse_text(document_path)
    raise DocumentParseError(f"没有适用于 {extension or '未知格式'} 的解析器")
