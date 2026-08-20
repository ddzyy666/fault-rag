from pathlib import Path

import pymupdf
from app.services.document_parser import parse_document
from docx import Document as WordDocument


def test_parse_pdf_by_page(tmp_path: Path) -> None:
    path = tmp_path / "manual.pdf"
    pdf = pymupdf.open()
    first_page = pdf.new_page()
    first_page.insert_text((72, 72), "Compressor alarm E101")
    second_page = pdf.new_page()
    second_page.insert_text((72, 72), "Check cooling fan")
    pdf.save(path)
    pdf.close()

    pages = parse_document(path)

    assert len(pages) == 2
    assert pages[0].page_number == 1
    assert "E101" in pages[0].content
    assert "cooling fan" in pages[1].content


def test_parse_docx_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "manual.docx"
    word = WordDocument()
    word.add_paragraph("故障现象：设备高温停机")
    table = word.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "原因"
    table.cell(0, 1).text = "散热器堵塞"
    word.save(path)

    pages = parse_document(path)

    assert len(pages) == 1
    assert "设备高温停机" in pages[0].content
    assert "散热器堵塞" in pages[0].content
